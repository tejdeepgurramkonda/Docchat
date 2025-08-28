"""
File Handler Module
Handles extraction of text from various file formats (PDF, DOCX, TXT)
"""

import os
import logging
from typing import Optional
import fitz  # PyMuPDF
from docx import Document

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class FileHandler:
    """
    Handles file operations and text extraction for various file formats
    """
    
    def __init__(self):
        self.supported_formats = ['pdf', 'docx', 'txt']
    
    def extract_text(self, file_path: str) -> str:
        """
        Extract text from a file based on its extension
        
        Args:
            file_path (str): Path to the file
            
        Returns:
            str: Extracted text content
            
        Raises:
            ValueError: If file format is not supported
            FileNotFoundError: If file doesn't exist
        """
        # Normalize the path for the current OS
        file_path = os.path.normpath(file_path)
        
        if not os.path.exists(file_path):
            # Try to find the file in the current directory
            current_dir = os.getcwd()
            alt_path = os.path.join(current_dir, file_path)
            if os.path.exists(alt_path):
                file_path = alt_path
            else:
                raise FileNotFoundError(f"File not found: {file_path}")
        
        logger.info(f"Processing file: {file_path}")
        
        file_extension = os.path.splitext(file_path)[1].lower().lstrip('.')
        
        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        try:
            if file_extension == 'pdf':
                return self._extract_from_pdf(file_path)
            elif file_extension == 'docx':
                return self._extract_from_docx(file_path)
            elif file_extension == 'txt':
                return self._extract_from_txt(file_path)
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            raise
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """
        Extract text from PDF file using PyMuPDF
        
        Args:
            file_path (str): Path to the PDF file
            
        Returns:
            str: Extracted text
        """
        try:
            doc = fitz.open(file_path)
            text = ""
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text += page.get_text()
            
            doc.close()
            
            if not text.strip():
                raise ValueError("No text found in PDF. The file might be image-based.")
            
            logger.info(f"Successfully extracted {len(text)} characters from PDF")
            return text
            
        except Exception as e:
            logger.error(f"Error reading PDF file: {str(e)}")
            raise
    
    def _extract_from_docx(self, file_path: str) -> str:
        """
        Extract text from DOCX file using python-docx
        
        Args:
            file_path (str): Path to the DOCX file
            
        Returns:
            str: Extracted text
        """
        try:
            doc = Document(file_path)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\t"
                    text += "\n"
            
            if not text.strip():
                raise ValueError("No text found in DOCX file.")
            
            logger.info(f"Successfully extracted {len(text)} characters from DOCX")
            return text
            
        except Exception as e:
            logger.error(f"Error reading DOCX file: {str(e)}")
            raise
    
    def _extract_from_txt(self, file_path: str) -> str:
        """
        Extract text from TXT file
        
        Args:
            file_path (str): Path to the TXT file
            
        Returns:
            str: Extracted text
        """
        try:
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                    
                    if not text.strip():
                        raise ValueError("The text file is empty.")
                    
                    logger.info(f"Successfully extracted {len(text)} characters from TXT")
                    return text
                    
                except UnicodeDecodeError:
                    continue
            
            raise ValueError("Unable to decode the text file with common encodings.")
            
        except Exception as e:
            logger.error(f"Error reading TXT file: {str(e)}")
            raise
    
    def get_file_info(self, file_path: str) -> dict:
        """
        Get information about a file
        
        Args:
            file_path (str): Path to the file
            
        Returns:
            dict: File information
        """
        try:
            stat = os.stat(file_path)
            return {
                'name': os.path.basename(file_path),
                'size': stat.st_size,
                'extension': os.path.splitext(file_path)[1].lower().lstrip('.'),
                'path': file_path
            }
        except Exception as e:
            logger.error(f"Error getting file info: {str(e)}")
            return {}
    
    def validate_file(self, file_path: str, max_size_mb: int = 100) -> bool:
        """
        Validate if a file can be processed
        
        Args:
            file_path (str): Path to the file
            max_size_mb (int): Maximum file size in MB
            
        Returns:
            bool: True if file is valid, False otherwise
        """
        try:
            if not os.path.exists(file_path):
                return False
            
            file_info = self.get_file_info(file_path)
            
            # Check file size
            if file_info['size'] > max_size_mb * 1024 * 1024:
                logger.error(f"File size ({file_info['size'] / 1024 / 1024:.1f} MB) exceeds maximum allowed size ({max_size_mb} MB)")
                return False
            
            # Check file extension
            if file_info['extension'] not in self.supported_formats:
                logger.error(f"Unsupported file format: {file_info['extension']}")
                return False
            
            return True
            
        except Exception as e:
            logger.error(f"Error validating file: {str(e)}")
            return False
